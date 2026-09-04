"""Export Service for Specifications."""
import logging
import io

try:
    from reportlab.pdfgen import canvas # type: ignore
    from reportlab.lib.pagesizes import letter # type: ignore
except ImportError:
    pass

try:
    import docx # type: ignore
except ImportError:
    pass

logger = logging.getLogger(__name__)

class ExportService:
    def export_pdf(self, spec: dict) -> bytes:
        buffer = io.BytesIO()
        try:
            c = canvas.Canvas(buffer, pagesize=letter)
            
            # Govt Header
            c.setFont("Helvetica-Bold", 14)
            c.drawCentredString(306, 760, "GOVERNMENT OF INDIA")
            c.setFont("Helvetica", 12)
            c.drawCentredString(306, 740, "STANDARD BIDDING DOCUMENT (SBD)")
            
            c.line(100, 725, 512, 725)
            
            c.setFont("Helvetica-Bold", 16)
            c.drawString(100, 700, str(spec.get('title', 'Specification')))
            
            y = 660
            sections = spec.get('sections', [])
            if not isinstance(sections, list):
                sections = []
                
            for section in sections:
                if y < 100:
                    c.showPage()
                    y = 750
                    
                heading = section.get('heading', '')
                content = section.get('content', '')
                
                c.setFont("Helvetica-Bold", 12)
                c.drawString(100, y, str(heading))
                y -= 20
                
                c.setFont("Helvetica", 10)
                # Simple text wrapping for PDF
                import textwrap
                wrapped_text = textwrap.wrap(str(content), width=85)
                for line in wrapped_text:
                    if y < 100:
                        c.showPage()
                        c.setFont("Helvetica", 10)
                        y = 750
                    c.drawString(100, y, line)
                    y -= 15
                y -= 20
                
            c.showPage()
            c.save()
        except Exception as e:
            logger.error(f"PDF export failed: {e}")
        return buffer.getvalue()

    def export_docx(self, spec: dict) -> bytes:
        buffer = io.BytesIO()
        try:
            doc = docx.Document()
            
            # Govt Header
            h1 = doc.add_paragraph()
            h1.alignment = 1 # Center
            r1 = h1.add_run("GOVERNMENT OF INDIA")
            r1.bold = True
            
            h2 = doc.add_paragraph()
            h2.alignment = 1
            r2 = h2.add_run("STANDARD BIDDING DOCUMENT (SBD)")
            
            doc.add_heading(str(spec.get('title', 'Specification')), 0)
            
            sections = spec.get('sections', [])
            if not isinstance(sections, list):
                sections = []
                
            for section in sections:
                doc.add_heading(str(section.get('heading', '')), level=1)
                doc.add_paragraph(str(section.get('content', '')))
                
            doc.save(buffer)
        except Exception as e:
            logger.error(f"DOCX export failed: {e}")
        return buffer.getvalue()
